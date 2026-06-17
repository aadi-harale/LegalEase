import asyncio
from fastapi import Depends, FastAPI, HTTPException, UploadFile, File, Request, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from pydantic import BaseModel
from datetime import datetime, timezone
from io import BytesIO
import os
import logging
import tempfile
import time
from typing import Optional
import uuid
import json

from dotenv import load_dotenv

from backend.database import engine, Base, SessionLocal, get_db
from sqlalchemy.orm import Session
from backend import models
from backend.routers import auth_routes
from backend.routers import legal_routes
from backend.routers import history_routes
from backend.auth import validate_token_or_api_key, AuthIdentity
from backend.utils.limiter import SimpleRateLimiter

# Optional imports (wrap in try/except so server can start without optional deps)
try:
    import fitz  # PyMuPDF
except Exception:
    fitz = None

try:
    from docx import Document as DocxDocument  # type: ignore[import-untyped]
except Exception:
    DocxDocument = None  # type: ignore[assignment,misc]

# Import pipeline exceptions, validations, and service
from backend.core.exceptions import (
    AIError, ValidationError, ProviderError, TimeoutError, ServiceUnavailableError
)
from backend.core.validation import (
    validate_chat_input, validate_summarize_input, sanitize_text, validate_mime_and_bytes,
    validate_docx_archive_safety
)
from backend.services.ai_service import ai_service, correlation_id_var

#Middleware import 
from backend.middleware.rate_limit import RateLimitMiddleware
from backend.middleware.correlation_id import validate_or_generate_correlation_id
# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv()

# Track application start time for uptime calculation
_app_start_time = time.monotonic()

app = FastAPI()


# Exception Handlers to match standardized error contracts
@app.exception_handler(ValidationError)
async def validation_exception_handler(request: Request, exc: ValidationError):
    logger.warning(f"[{correlation_id_var.get()}] Validation error: {exc}")
    return JSONResponse(
        status_code=400,
        content={
            "error": "validation_error",
            "detail": str(exc),
            "correlation_id": correlation_id_var.get()
        }
    )


@app.exception_handler(ProviderError)
async def provider_exception_handler(request: Request, exc: ProviderError):
    logger.error(f"[{correlation_id_var.get()}] Upstream provider error: {exc}")
    return JSONResponse(
        status_code=502,
        content={
            "error": "provider_error",
            "detail": str(exc),
            "correlation_id": correlation_id_var.get()
        }
    )


@app.exception_handler(TimeoutError)
async def timeout_exception_handler(request: Request, exc: TimeoutError):
    logger.error(f"[{correlation_id_var.get()}] Request timeout: {exc}")
    return JSONResponse(
        status_code=504,
        content={
            "error": "timeout_error",
            "detail": str(exc),
            "correlation_id": correlation_id_var.get()
        }
    )


@app.exception_handler(ServiceUnavailableError)
async def service_unavailable_exception_handler(request: Request, exc: ServiceUnavailableError):
    logger.error(f"[{correlation_id_var.get()}] Service unavailable: {exc}")
    return JSONResponse(
        status_code=503,
        content={
            "error": "service_unavailable",
            "detail": str(exc),
            "correlation_id": correlation_id_var.get()
        }
    )

import sys

# Global unhandled HTTP exceptions
@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    corr_id = correlation_id_var.get()
    logger.error(f"[{corr_id}] Unhandled exception: {exc}", exc_info=True)
    return JSONResponse(
        status_code=500,
        content={
            "error": "internal_server_error",
            "detail": "An unexpected error occurred.",
            "correlation_id": corr_id
        }
    )

# Global unhandled thread/process exceptions
def handle_uncaught_exception(exc_type, exc_value, exc_traceback):
    if not issubclass(exc_type, Exception):
        sys.__excepthook__(exc_type, exc_value, exc_traceback)
        return
    logger.critical("Uncaught global exception", exc_info=(exc_type, exc_value, exc_traceback))

sys.excepthook = handle_uncaught_exception


# Create database tables
Base.metadata.create_all(bind=engine)

# Include authentication router
app.include_router(auth_routes.router)
# Include legal mapping router
app.include_router(legal_routes.router)
# Include history router
app.include_router(history_routes.router)


# Enable CORS for frontend communication
raw_allowed_origins = os.getenv("ALLOWED_ORIGINS") or os.getenv(
    "FRONTEND_URL",
    "http://localhost:5173"
)
ALLOWED_ORIGINS = [
    origin.strip()
    for origin in raw_allowed_origins.split(",")
    if origin.strip()
]
# Automatically allow common development ports on localhost
for host in ["http://localhost", "http://127.0.0.1"]:
    for port in range(5173, 5181):
        dev_origin = f"{host}:{port}"
        if dev_origin not in ALLOWED_ORIGINS:
            ALLOWED_ORIGINS.append(dev_origin)
# Rate-limit middleware registered first so that CORSMiddleware
# (added second) wraps it — ensuring 429 responses include CORS headers.
app.add_middleware(RateLimitMiddleware)
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
logger.info(f"Allowed frontend origins: {ALLOWED_ORIGINS}")


# Correlation ID middleware to inject trace headers
@app.middleware("http")
async def correlation_id_middleware(request: Request, call_next):
    # Validate or generate safe correlation ID
    client_id = request.headers.get("X-Correlation-ID")
    corr_id, was_valid = validate_or_generate_correlation_id(client_id)
    
    token = correlation_id_var.set(corr_id)
    try:
        response = await call_next(request)
        response.headers["X-Correlation-ID"] = corr_id
        return response
    finally:
        correlation_id_var.reset(token)


# Configuration
MAX_UPLOAD_SIZE = int(os.getenv("MAX_UPLOAD_SIZE", str(25 * 1024 * 1024)))  # 25 MB default
CHUNK_SIZE = 1024 * 1024
MAX_PDF_PAGES = int(os.getenv("MAX_PDF_PAGES", "100"))
MAX_DOCX_PARAGRAPHS = int(os.getenv("MAX_DOCX_PARAGRAPHS", "2000"))
MAX_EXTRACTED_TEXT_CHARS = int(os.getenv("MAX_EXTRACTED_TEXT_CHARS", "10000"))
UPLOAD_PARSE_TIMEOUT_SECONDS = float(os.getenv("UPLOAD_PARSE_TIMEOUT_SECONDS", "5"))


RATE_LIMIT_PERIOD = int(os.getenv("RATE_LIMIT_PERIOD", "60"))
RATE_LIMIT_KEY_CALLS = int(os.getenv("RATE_LIMIT_KEY_CALLS", "300"))


# Defaults: 300 requests per minute per API key
key_limiter = SimpleRateLimiter(calls=RATE_LIMIT_KEY_CALLS, period=RATE_LIMIT_PERIOD)


class ChatRequest(BaseModel):
    message: str
    context: Optional[str] = None
    conversation_history: Optional[list[dict[str, str]]] = None
    stream: Optional[bool] = False


class SummarizeRequest(BaseModel):
    text: str


class HealthResponse(BaseModel):
    status: str
    uptime_seconds: float
    timestamp: str
    details: Optional[dict] = None


def _validate_api_key(request: Request) -> str:
    # Accept header `Authorization: Bearer <key>` or `X-API-Key`
    auth = request.headers.get("authorization") or ""
    api_key = ""
    if auth.lower().startswith("bearer "):
        api_key = auth.split(" ", 1)[1].strip()
    else:
        api_key = request.headers.get("x-api-key", "").strip()

    if not api_key:
        raise HTTPException(status_code=401, detail="Missing API key")

    # Read from environment dynamically (allows test mocking)
    api_keys = [k.strip() for k in os.getenv("API_KEYS", "").split(",") if k.strip()]
    allow_dev = os.getenv("ALLOW_DEV", "false").lower() in ("1", "true", "yes")
    dev_api_key = os.getenv("DEV_API_KEY", "dev-token")

    # Check production API keys first
    if api_key in api_keys:
        return api_key

    # Check dev mode only when production keys are not configured
    if not api_keys and allow_dev and api_key == dev_api_key:
        return api_key

    raise HTTPException(status_code=403, detail="Invalid API key")

def _get_client_ip(request: Request) -> str:
    forwarded_for = request.headers.get("x-forwarded-for")
    if forwarded_for:
        return forwarded_for.split(",", 1)[0].strip()
    return request.client.host if request.client else "unknown"


def _extract_pdf_text(file_path: str) -> str:
    if fitz is None:
        raise HTTPException(status_code=503, detail="PDF processing not available")

    doc = None
    try:
        doc = fitz.open(file_path)
        if doc.page_count > MAX_PDF_PAGES:
            raise HTTPException(status_code=413, detail="PDF is too large to process safely")

        extracted_parts = []
        for page in doc:
            extracted_parts.append(page.get_text())
            if len("".join(extracted_parts)) >= MAX_EXTRACTED_TEXT_CHARS:
                break

        return "".join(extracted_parts)
    finally:
        if doc is not None:
            doc.close()


def _extract_docx_text(file_path: str) -> str:
    if DocxDocument is None:
        raise HTTPException(status_code=503, detail="DOCX processing not available")

    document = None
    try:
        document = DocxDocument(file_path)
        if len(document.paragraphs) > MAX_DOCX_PARAGRAPHS:
            raise HTTPException(status_code=413, detail="DOCX is too large to process safely")

        return "\n".join(
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        )
    finally:
        close_method = getattr(document, "close", None) if document is not None else None
        if callable(close_method):
            close_method()


async def _run_bounded_parser(parser, file_path: str) -> str:
    try:
        return await asyncio.wait_for(
            asyncio.to_thread(parser, file_path),
            timeout=UPLOAD_PARSE_TIMEOUT_SECONDS,
        )
    except asyncio.TimeoutError:
        raise HTTPException(status_code=413, detail="File is too complex to process safely")

async def analyze_document_background(doc_id: int, text: str):
    db = SessionLocal()
    try:
        from backend.services.ai_service import ai_service
        result = await ai_service.analyze_clauses(text)
        
        doc = db.query(models.DocumentRecord).filter(models.DocumentRecord.id == doc_id).first()
        if doc:
            doc.liability_score = result.get("liabilityScore", 50)
            doc.risk_analysis = json.dumps(result.get("clauses", []))
            doc.status = "completed"
            db.commit()
    except Exception as e:
        logger.error(f"Background analysis failed for doc {doc_id}: {e}")
        doc = db.query(models.DocumentRecord).filter(models.DocumentRecord.id == doc_id).first()
        if doc:
            doc.status = "failed"
            db.commit()
    finally:
        db.close()


@app.post("/chat")
async def chat(request: Request, payload: ChatRequest, identity: AuthIdentity = Depends(validate_token_or_api_key)):
    # Rate limiting using the authenticated identity
    if not key_limiter.check(identity.get_rate_limit_key())["allowed"]:
        raise HTTPException(status_code=429, detail="Rate limit exceeded")

    # Sanitize inputs
    sanitized_message = sanitize_text(payload.message)
    sanitized_context = sanitize_text(payload.context) if payload.context else None

    # Early payload validation
    validate_chat_input(sanitized_message, sanitized_context)

    # Streaming or standard block handling
    if payload.stream:
        async def stream_generator():
            try:
                async for chunk in ai_service.generate_chat_response(
                    message=sanitized_message,
                    context=sanitized_context,
                    history=payload.conversation_history,
                    stream=True
                ):
                    yield chunk
            except Exception as e:
                logger.error(f"[{correlation_id_var.get()}] Stream generation error: {e}")
                yield "\n[Error: Inference stream failed]"

        return StreamingResponse(stream_generator(), media_type="text/event-stream")
    else:
        response_gen = ai_service.generate_chat_response(
            message=sanitized_message,
            context=sanitized_context,
            history=payload.conversation_history,
            stream=False
        )
        response_text = ""
        async for chunk in response_gen:
            response_text += chunk
        return {"response": response_text}


@app.post("/upload")
async def upload_document(
    request: Request,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    identity: AuthIdentity = Depends(validate_token_or_api_key),
    db: Session = Depends(get_db)
):
    # Content-Length pre-check
    try:
        content_length = int(request.headers.get("content-length", "0"))
    except Exception:
        content_length = 0
    if content_length and content_length > MAX_UPLOAD_SIZE:
        raise HTTPException(status_code=413, detail="Uploaded file is too large")

    temp_path = None
    try:
        filename = file.filename or "unknown"
        file_extension = os.path.splitext(filename)[1].lower()
        extracted_text = ""
        total_size = 0
        with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension or "") as temp_file:
            temp_path = temp_file.name
            while True:
                chunk = await file.read(CHUNK_SIZE)
                if not chunk:
                    break
                total_size += len(chunk)
                if total_size > MAX_UPLOAD_SIZE:
                    raise HTTPException(status_code=413, detail="Uploaded file is too large")
                temp_file.write(chunk)

        # Read only first 4096 bytes for MIME validation
        with open(temp_path, "rb") as temp_file:
            content_prefix = temp_file.read(4096)

        # Perform MIME-aware preprocessing and signature validation using only the
        # minimum bytes needed for structural checks.
        validate_mime_and_bytes(content_prefix, file.content_type or "", filename)

        if file_extension == ".docx":
            validate_docx_archive_safety(temp_path)

        # Process by type using file path to avoid loading entire file into memory
        if file_extension == '.pdf' or content_prefix.startswith(b'%PDF-'):
            try:
                extracted_text = await _run_bounded_parser(_extract_pdf_text, temp_path)
            except HTTPException:
                raise
            except Exception as e:
                if isinstance(e, HTTPException):
                    raise
                logger.error(f"PDF parse error: {e}")
                raise HTTPException(status_code=400, detail="Invalid or corrupted PDF")

        elif file_extension == '.docx':
            try:
                extracted_text = await _run_bounded_parser(_extract_docx_text, temp_path)
            except HTTPException:
                raise
            except Exception:
                raise HTTPException(
                    status_code=400,
                    detail="Invalid or corrupted DOCX file."
                )

        elif file_extension == '.txt':
            with open(temp_path, 'r', encoding='utf-8') as text_file:
                extracted_text = text_file.read(10000)

        # Truncate extracted text to avoid sending huge payloads to models
        extracted_text = extracted_text[:MAX_EXTRACTED_TEXT_CHARS]

        doc_id = None
        user_id = identity.get_user_id()
        if user_id:
            new_doc = models.DocumentRecord(
                user_id=user_id,
                filename=filename,
                file_type=file.content_type or "",
                summary="",
                status="processing"
            )
            db.add(new_doc)
            db.commit()
            db.refresh(new_doc)
            doc_id = new_doc.id
            
            # Queue background task
            background_tasks.add_task(analyze_document_background, doc_id, extracted_text)

        return {"filename": filename, "text": extracted_text, "document_id": doc_id}

    except ValidationError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Upload error: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Failed to process document")
    finally:
        if temp_path and os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
            except OSError:
                pass


@app.post("/summarize")
async def summarize(request: Request, payload: SummarizeRequest, identity: AuthIdentity = Depends(validate_token_or_api_key)):

    # Sanitize input
    sanitized_text = sanitize_text(payload.text)

    # Early payload validation
    validate_summarize_input(sanitized_text)

    summary = await ai_service.generate_summary(sanitized_text)
    return {"summary": summary}


@app.get("/health", response_model=HealthResponse)
async def health():
    """
    Health check endpoint with structured response.
    Returns HTTP 503 when the service is degraded.
    """
    health_data = ai_service.check_health()
    uptime = time.monotonic() - _app_start_time
    timestamp = datetime.now(timezone.utc).isoformat().replace("+00:00", "Z")

    # Database connectivity check
    db_status = "ok"
    try:
        from sqlalchemy import text
        db = SessionLocal()
        db.execute(text("SELECT 1"))
        db.close()
    except Exception as e:
        logger.error(f"Database health check failed: {e}")
        db_status = "down"

    status = health_data.get("status", "unknown")
    if db_status == "down":
        status = "degraded"

    details = health_data.get("details") or {}
    if not isinstance(details, dict):
        details = {"ai_details": details}
    details["database"] = db_status

    response = HealthResponse(
        status=status,
        uptime_seconds=round(uptime, 2),
        timestamp=timestamp,
        details=details,
    )

    if response.status == "degraded":
        return JSONResponse(status_code=503, content={"detail": response.model_dump()})

    return response


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)
